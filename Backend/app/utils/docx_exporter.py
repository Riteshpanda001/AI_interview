import io
import zipfile
from typing import Dict, Any

class DOCXExporter:
    """
    Utility class to convert structured resume JSON into a styled Word .docx document.
    Uses python-docx when available, with a clean OpenXML fallback.
    """

    @staticmethod
    def generate_docx(resume_data: Dict[str, Any], title: str = "Resume") -> bytes:
        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            doc = docx.Document()

            # Set margins (0.5 inch)
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.6)
                section.right_margin = Inches(0.6)

            personal = resume_data.get("personal", {})
            name = personal.get("name", "Candidate Name")
            role = personal.get("role", "")
            email = personal.get("email", "")
            phone = personal.get("phone", "")
            location = personal.get("location", "")
            linkedin = personal.get("linkedin", "")
            github = personal.get("github", "")

            # Title / Header
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_name = header_p.add_run(name.upper())
            run_name.font.size = Pt(22)
            run_name.font.bold = True
            run_name.font.color.rgb = RGBColor(26, 43, 76)

            if role:
                header_p.add_run(f"\n{role}").font.size = Pt(13)

            contact_parts = [p for p in [phone, email, location, linkedin, github] if p]
            if contact_parts:
                contact_p = doc.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_contact = contact_p.add_run(" | ".join(contact_parts))
                run_contact.font.size = Pt(9.5)
                run_contact.font.color.rgb = RGBColor(100, 100, 100)

            def add_section_header(section_title: str):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(section_title.upper())
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 43, 76)

            # Professional Summary
            summary = resume_data.get("summary")
            if summary:
                add_section_header("Professional Summary")
                sp = doc.add_paragraph(summary)
                sp.paragraph_format.space_after = Pt(6)
                sp.runs[0].font.size = Pt(10)

            # Technical Skills
            skills = resume_data.get("skills", [])
            if skills:
                add_section_header("Skills & Expertise")
                skills_str = ", ".join(skills) if isinstance(skills, list) else str(skills)
                sk_p = doc.add_paragraph(skills_str)
                sk_p.paragraph_format.space_after = Pt(6)
                sk_p.runs[0].font.size = Pt(10)

            # Work Experience
            experience = resume_data.get("experience", [])
            if experience and isinstance(experience, list):
                add_section_header("Work Experience")
                for exp in experience:
                    comp = exp.get("company", "")
                    pos = exp.get("role", "")
                    dur = exp.get("duration", "")
                    loc = exp.get("location", "")

                    exp_p = doc.add_paragraph()
                    exp_p.paragraph_format.space_before = Pt(4)
                    exp_p.paragraph_format.space_after = Pt(2)
                    r_pos = exp_p.add_run(pos)
                    r_pos.bold = True
                    r_pos.font.size = Pt(10.5)

                    if comp:
                        r_comp = exp_p.add_run(f" — {comp}")
                        r_comp.font.size = Pt(10.5)
                    if dur:
                        r_dur = exp_p.add_run(f"\t{dur}")
                        r_dur.font.size = Pt(9.5)
                        r_dur.font.color.rgb = RGBColor(100, 100, 100)

                    details = exp.get("details", "")
                    if details:
                        lines = [l.strip("•- ").strip() for l in details.split("\n") if l.strip()]
                        for line in lines:
                            bp = doc.add_paragraph(style='List Bullet')
                            bp.paragraph_format.space_after = Pt(2)
                            r_b = bp.add_run(line)
                            r_b.font.size = Pt(9.5)

            # Projects
            projects = resume_data.get("projects", [])
            if projects and isinstance(projects, list):
                add_section_header("Projects")
                for proj in projects:
                    pname = proj.get("name", "")
                    pdesc = proj.get("description", "")
                    ptech = proj.get("technologies", "")

                    pp = doc.add_paragraph()
                    pp.paragraph_format.space_before = Pt(4)
                    pp.paragraph_format.space_after = Pt(2)
                    r_pn = pp.add_run(pname)
                    r_pn.bold = True
                    r_pn.font.size = Pt(10.5)

                    if ptech:
                        r_pt = pp.add_run(f" ({ptech})")
                        r_pt.font.size = Pt(9.5)
                        r_pt.font.color.rgb = RGBColor(100, 100, 100)

                    if pdesc:
                        pdp = doc.add_paragraph(style='List Bullet')
                        pdp.paragraph_format.space_after = Pt(2)
                        pdp.add_run(pdesc).font.size = Pt(9.5)

            # Education
            education = resume_data.get("education", [])
            if education and isinstance(education, list):
                add_section_header("Education")
                for edu in education:
                    inst = edu.get("institution", "") or edu.get("school", "")
                    deg = edu.get("degree", "")
                    dur = edu.get("duration", "") or edu.get("year", "")

                    ep = doc.add_paragraph()
                    ep.paragraph_format.space_after = Pt(3)
                    r_deg = ep.add_run(deg)
                    r_deg.bold = True
                    r_deg.font.size = Pt(10)
                    if inst:
                        ep.add_run(f", {inst}").font.size = Pt(10)
                    if dur:
                        r_d = ep.add_run(f"\t{dur}")
                        r_d.font.size = Pt(9.5)
                        r_d.font.color.rgb = RGBColor(100, 100, 100)

            # Certifications
            certs = resume_data.get("certifications", [])
            if certs and isinstance(certs, list):
                add_section_header("Certifications")
                for cert in certs:
                    ctitle = cert.get("title", "") if isinstance(cert, dict) else str(cert)
                    cissuer = cert.get("issuer", "") if isinstance(cert, dict) else ""
                    cyear = cert.get("year", "") if isinstance(cert, dict) else ""

                    cp = doc.add_paragraph(style='List Bullet')
                    cp.paragraph_format.space_after = Pt(2)
                    r_ct = cp.add_run(ctitle)
                    r_ct.bold = True
                    r_ct.font.size = Pt(9.5)
                    if cissuer:
                        cp.add_run(f" - {cissuer}").font.size = Pt(9.5)
                    if cyear:
                        cp.add_run(f" ({cyear})").font.size = Pt(9.5)

            out_stream = io.BytesIO()
            doc.save(out_stream)
            return out_stream.getvalue()

        except Exception as e:
            print(f"python-docx error, fallback docx generation: {e}")
            return DOCXExporter._generate_openxml_docx_fallback(resume_data, title)

    @staticmethod
    def _generate_openxml_docx_fallback(resume_data: Dict[str, Any], title: str) -> bytes:
        """
        Creates a minimal valid OpenXML .docx ZIP container as a fail-safe fallback.
        """
        personal = resume_data.get("personal", {})
        name = personal.get("name", "Candidate")
        summary = resume_data.get("summary", "")

        doc_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="36"/></w:rPr><w:t>{name}</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:i/><w:sz w:val="24"/></w:rPr><w:t>{personal.get('role', '')}</w:t></w:r></w:p>
    <w:p><w:r><w:t>{summary}</w:t></w:r></w:p>
  </w:body>
</w:document>"""

        content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

        rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as z:
            z.writestr('[Content_Types].xml', content_types)
            z.writestr('_rels/.rels', rels)
            z.writestr('word/document.xml', doc_xml)

        return buffer.getvalue()
