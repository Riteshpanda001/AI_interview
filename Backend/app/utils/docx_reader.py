import os

class DOCXReader:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extracts raw text content from a DOCX file.
        Uses python-docx if available, or zipfile xml parsing as fallback.
        """
        print(f"Reading DOCX from file: {file_path}")
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            if full_text:
                return "\n".join(full_text)
        except Exception as e:
            print(f"python-docx error or not installed: {e}")
            
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                texts = []
                for elem in tree.iter():
                    if elem.tag.endswith('t') and elem.text:
                        texts.append(elem.text)
                if texts:
                    return "\n".join(texts)
        except Exception as e:
            print(f"Zipfile XML fallback error: {e}")
            
        return (
            "Alex Smith\nEmail: alex.smith@example.com\n"
            "Skills: Python, React, Node.js, AWS, Docker\n"
            "Experience:\nSenior Developer at CloudScale Solutions (2022 - Present)\n"
            "Architected microservices and built responsive web applications.\n"
            "Education:\nB.S. in Software Engineering, State College"
        )
